"""
LaaS Portal — End User Manual Generator
Generates: USER_MANUAL.docx
Run: python generate_manual_docx.py
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy, os

OUT = r"c:\caas-dashboard\USER_MANUAL.docx"

# ── colour palette ──────────────────────────────────────────────────────────
NAVY    = RGBColor(0x0d, 0x18, 0x29)
ACCENT  = RGBColor(0x3b, 0x82, 0xf6)
CYAN    = RGBColor(0x06, 0xb6, 0xd4)
GREEN   = RGBColor(0x10, 0xb9, 0x81)
ORANGE  = RGBColor(0xf9, 0x73, 0x16)
RED     = RGBColor(0xef, 0x44, 0x44)
PURPLE  = RGBColor(0x8b, 0x5c, 0xf6)
YELLOW  = RGBColor(0xf5, 0x9e, 0x0b)
WHITE   = RGBColor(0xff, 0xff, 0xff)
LIGHT   = RGBColor(0xe2, 0xe8, 0xf0)
SUB     = RGBColor(0x94, 0xa3, 0xb8)
MUTE    = RGBColor(0x47, 0x55, 0x69)
DARK_BG = RGBColor(0x1e, 0x29, 0x3b)
MID_BG  = RGBColor(0x0f, 0x17, 0x2a)


def shade_cell(cell, rgb: RGBColor):
    """Fill a table cell with a solid background colour."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    hex_color = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)


def add_page_break(doc):
    doc.add_page_break()


def heading(doc, text, level=1, color=None):
    h = doc.add_heading(text, level=level)
    h.alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in h.runs:
        run.font.color.rgb = color or ACCENT
        if level == 1:
            run.font.size = Pt(22)
            run.bold = True
        elif level == 2:
            run.font.size = Pt(16)
            run.bold = True
        else:
            run.font.size = Pt(13)
            run.bold = True
    return h


def para(doc, text, color=None, bold=False, size=11, indent=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.left_indent = Cm(indent)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color or SUB
    run.bold = bold
    return p


def bullet(doc, text, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(11)
    run.font.color.rgb = color or SUB


def divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '1E293B')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_table(doc, headers, rows, col_widths=None, header_color=None):
    n_cols = len(headers)
    table = doc.add_table(rows=1 + len(rows), cols=n_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        shade_cell(cell, header_color or DARK_BG)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        for run in p.runs:
            run.font.bold = True
            run.font.size = Pt(10)
            run.font.color.rgb = ACCENT

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        bg = RGBColor(0x0d, 0x18, 0x29) if r_idx % 2 == 0 else RGBColor(0x11, 0x18, 0x27)
        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(cell_text)
            shade_cell(cell, bg)
            p = cell.paragraphs[0]
            for run in p.runs:
                run.font.size = Pt(10)
                run.font.color.rgb = SUB

    if col_widths:
        for i, w in enumerate(col_widths):
            set_col_width(table, i, w)

    doc.add_paragraph()  # spacing
    return table


def cover_page(doc):
    # Title block
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(60)
    run = p.add_run("⚡  LaaS Portal")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("End User Manual & Technical Reference Guide")
    r2.font.size = Pt(16)
    r2.font.color.rgb = SUB

    doc.add_paragraph()

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run("Lab as a Service  ·  Unified Infrastructure Management Platform")
    r3.font.size = Pt(13)
    r3.font.color.rgb = CYAN
    r3.italic = True

    doc.add_paragraph()
    divider(doc)
    doc.add_paragraph()

    # Meta table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Version", "6.0"),
        ("Date", "March 2026"),
        ("Organisation", "Wipro SDX INFRA & DC Operations"),
        ("Prepared By", "Sekhar Perumal"),
        ("Approved By", "Khalid Khan"),
    ]
    for i, (k, v) in enumerate(meta_data):
        row = meta_table.rows[i]
        bg = RGBColor(0x1e, 0x29, 0x3b)
        shade_cell(row.cells[0], bg)
        shade_cell(row.cells[1], bg)
        row.cells[0].text = k
        row.cells[1].text = v
        for cell, clr in [(row.cells[0], MUTE), (row.cells[1], LIGHT)]:
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(11)
                run.font.color.rgb = clr
                if cell == row.cells[0]:
                    run.bold = True
        set_col_width(meta_table, 0, 5)
        set_col_width(meta_table, 1, 9)

    doc.add_paragraph()

    # Platforms
    p_plat = doc.add_paragraph()
    p_plat.alignment = WD_ALIGN_PARAGRAPH.CENTER
    platforms = ["🖥️ VMware vCenter", "🔴 Red Hat OpenShift", "🟦 Nutanix AHV", "🎛️ Ansible AAP", "📡 IPAM"]
    for i, plt in enumerate(platforms):
        run = p_plat.add_run(plt + ("   ·   " if i < len(platforms)-1 else ""))
        run.font.size = Pt(11)
        run.font.color.rgb = CYAN

    add_page_break(doc)


def toc_page(doc):
    heading(doc, "Table of Contents", 1, ACCENT)
    divider(doc)
    doc.add_paragraph()

    toc_items = [
        ("01", "Introduction & About LaaS Portal"),
        ("02", "System Requirements"),
        ("03", "Quick Start Guide"),
        ("04", "Navigation & UI Overview"),
        ("05", "Overview Dashboard"),
        ("06", "VMware Module"),
        ("07", "Snapshot Management"),
        ("08", "Red Hat OpenShift Module"),
        ("09", "Nutanix AHV Module"),
        ("10", "Ansible AAP Module"),
        ("11", "Capacity Planning"),
        ("12", "Project Utilization"),
        ("13", "Chargeback"),
        ("14", "Networks & IPAM"),
        ("15", "VM Request Workflow"),
        ("16", "Audit Log"),
        ("17", "User Roles & Permissions"),
        ("18", "Technical Specifications"),
        ("19", "FAQ & Troubleshooting"),
        ("20", "Document Authorization"),
    ]
    for num, label in toc_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"  {num}  ")
        r1.font.size = Pt(11)
        r1.font.color.rgb = ACCENT
        r1.bold = True
        r2 = p.add_run(label)
        r2.font.size = Pt(11)
        r2.font.color.rgb = SUB

    add_page_break(doc)


def section_intro(doc):
    heading(doc, "1.  Introduction & About LaaS Portal", 1)
    divider(doc)
    para(doc,
         "The LaaS Portal (Lab as a Service) is Wipro's unified infrastructure management platform for "
         "SDX INFRA & DC Operations. It provides a single-pane-of-glass web dashboard to manage, monitor, "
         "and operate multi-cloud and on-premises resources across VMware vCenter, Red Hat OpenShift, "
         "Nutanix AHV, and Ansible Automation Platform.",
         color=SUB)
    para(doc,
         "The portal eliminates the need to log in to individual vCenter/OCP/Nutanix consoles for routine "
         "operations. It consolidates health monitoring, VM lifecycle management, capacity reporting, IP address "
         "management, chargeback, and audit logging into a single role-based access-controlled interface "
         "accessible from any modern browser.",
         color=SUB)

    heading(doc, "Key Features", 2, CYAN)
    features = [
        ("Unified Multi-Platform View", "Single dashboard for VMware, OpenShift, Nutanix, and Ansible — no separate logins."),
        ("VM Lifecycle Management", "Power on/off, clone, migrate, snapshot, and delete VMs directly from the portal."),
        ("Capacity Planning", "CPU, RAM, and storage utilisation trends with forecasting gauges."),
        ("Role-Based Access Control", "Admin, Operator, Viewer, and Requester roles with granular permissions."),
        ("VM Request & Approval Workflow", "Requesters raise VM provisioning requests; Admins approve or reject."),
        ("Automated Chargeback", "Project-based cost allocation with monthly PDF/CSV export."),
        ("IPAM Integration", "Live IP address management — subnet occupancy, free IPs, allocation history."),
        ("Full Audit Log", "Every action logged with user, timestamp, and resource for compliance."),
    ]
    for title, desc in features:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"{title}:  ")
        r1.bold = True
        r1.font.color.rgb = LIGHT
        r1.font.size = Pt(11)
        r2 = p.add_run(desc)
        r2.font.size = Pt(11)
        r2.font.color.rgb = SUB

    add_page_break(doc)


def section_sysreq(doc):
    heading(doc, "2.  System Requirements", 1)
    divider(doc)
    para(doc, "End-user browser and network requirements for accessing the LaaS Portal:", color=SUB)
    add_table(doc,
        ["Requirement", "Details"],
        [
            ("Supported Browsers", "Google Chrome 110+, Microsoft Edge 110+, Firefox 115+, Safari 16+"),
            ("JavaScript", "Must be enabled (portal is a React single-page application)"),
            ("Screen Resolution", "Minimum 1280×768 · Recommended 1920×1080 or higher"),
            ("Network", "Corporate LAN or VPN connection to the portal server required"),
            ("Portal URL", "http://<server-ip>:5173 (dev) · http://<server-ip>:80 (production)"),
            ("Authentication", "Username + Password (local accounts managed by Admin)"),
            ("Session Timeout", "Auto-logout after 30 minutes of inactivity"),
        ],
        col_widths=[6, 12]
    )
    add_page_break(doc)


def section_quickstart(doc):
    heading(doc, "3.  Quick Start Guide", 1)
    divider(doc)
    steps = [
        ("Open the Portal",
         "Navigate to http://<server-ip> in your browser. Ensure you are connected to the corporate network or VPN."),
        ("Log In",
         "Enter your username and password. If you don't have credentials, contact your Admin.\n"
         "Default demo credentials: admin / admin123"),
        ("Select Environment",
         "Use the vCenter dropdown in the top bar to choose a specific vCenter or 'All vCenters' for an aggregated view."),
        ("Explore the Dashboard",
         "The Overview page shows your SDx Infrastructure banner with live counts for VMs, hosts, OCP nodes, pods, and Nutanix resources."),
        ("Navigate by Platform",
         "Use the left sidebar to switch between VMware, OpenShift, Nutanix, Ansible, IPAM, and other modules."),
        ("Global Search",
         "Press Ctrl+K or use the search bar to instantly jump to any VM, snapshot, OCP cluster, or menu item."),
    ]
    for i, (title, body) in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        r1 = p.add_run(f"  Step {i}:  ")
        r1.bold = True
        r1.font.color.rgb = ACCENT
        r1.font.size = Pt(12)
        r2 = p.add_run(title)
        r2.bold = True
        r2.font.color.rgb = LIGHT
        r2.font.size = Pt(12)
        para(doc, body, color=SUB, indent=1.5)

    add_page_break(doc)


def section_navigation(doc):
    heading(doc, "4.  Navigation & UI Overview", 1)
    divider(doc)
    para(doc,
         "The portal is divided into three zones: Left Sidebar (navigation + clock + health status), "
         "Top Bar (vCenter selector + search + user profile), and Main Content Area (page content).",
         color=SUB)

    heading(doc, "Sidebar Navigation Items", 2, CYAN)
    add_table(doc,
        ["Icon", "Label", "Description", "Roles"],
        [
            ("🏠", "Overview", "SDx Infrastructure banner, platform health, KPI cards", "All"),
            ("🖥️", "VMware", "VM list, actions (power/clone/migrate/snapshot)", "Admin, Operator, Viewer"),
            ("  📸", "↳ Snapshots", "All VM snapshots — nested under VMware", "Admin, Operator, Viewer"),
            ("🔴", "Red Hat OpenShift", "Cluster status, node/pod/operator health", "Admin, Operator, Viewer"),
            ("🟦", "Nutanix", "Prism Central overview, cluster/VM/storage health", "Admin, Operator, Viewer"),
            ("🎛️", "Ansible AAP", "Job templates, inventory, recent job runs", "Admin, Operator, Viewer"),
            ("🏷️", "Project Utilization", "Per-project VM, CPU, RAM, storage allocation", "Admin, Operator, Viewer"),
            ("🗄️", "Capacity", "CPU/RAM/storage gauges, trend charts, datastores", "Admin, Operator, Viewer"),
            ("💳", "Chargeback", "Monthly cost by project, export to PDF/CSV", "Admin, Operator, Viewer"),
            ("🌐", "Networks", "vCenter port groups, VLANs, connected VMs", "Admin, Operator, Viewer"),
            ("📡", "IPAM", "Subnet overview, IP allocation table, free IP finder", "Admin, Operator, Viewer"),
            ("🗃️", "Asset Mgmt", "Physical and virtual asset inventory", "Admin, Operator, Viewer"),
            ("📋", "VM Requests", "Raise new VM request, view pending/approved/rejected", "All"),
            ("🔍", "Audit Log", "Complete action history with filters by user/date/type", "Admin, Operator"),
            ("🌲", "AD & DNS", "Active Directory and DNS record management", "Admin, Operator"),
            ("👥", "Users & Roles", "Create/edit users, assign roles", "Admin only"),
        ],
        col_widths=[1.2, 3.5, 7, 4]
    )
    add_page_break(doc)


def section_overview(doc):
    heading(doc, "5.  Overview Dashboard", 1)
    divider(doc)
    heading(doc, "SDx Infrastructure Overview Banner", 2, CYAN)
    para(doc,
         "The large banner at the top of the Overview page shows live aggregate statistics across all managed platforms.",
         color=SUB)

    bullet(doc, "Platform Pills Row: VMware vCenter count · Red Hat OpenShift cluster count · Nutanix Prism Central count — each with a live health indicator (healthy/degraded/critical).")
    bullet(doc, "Total VMs — sum of VMware + Nutanix VMs")
    bullet(doc, "ESXi Hosts — total connected ESXi hosts")
    bullet(doc, "OCP Nodes — total OpenShift nodes with ready count")
    bullet(doc, "Nutanix Hosts — AHV hosts with managed VM count")
    bullet(doc, "OCP Pods — running container pods")

    heading(doc, "Below the Banner", 2, CYAN)
    bullet(doc, "Environment Charts Panel — side-by-side bar/gauge charts for VMware CPU/RAM, OCP cluster utilisation, Nutanix storage.")
    bullet(doc, "vCenter Summary Cards — per-vCenter CPU/RAM/Storage bars with VM/Host counts.")
    bullet(doc, "Quick KPI Cards — Running VMs, ESXi Hosts, Active Alerts, Storage Free.")
    bullet(doc, "Capacity Detail Cards — full resource rows with utilisation bars.")

    add_page_break(doc)


def section_vmware(doc):
    heading(doc, "6.  VMware Module", 1)
    divider(doc)
    para(doc, "The VMware module provides full VM lifecycle management across all registered vCenter instances.", color=SUB)

    heading(doc, "VM Table Columns", 2, CYAN)
    add_table(doc,
        ["Column", "Description"],
        [
            ("VM Name", "Display name of the virtual machine"),
            ("Status", "Power state: Running / Stopped / Suspended"),
            ("IP Address", "Primary guest IP (requires VMware Tools installed)"),
            ("OS", "Guest OS type detected via VMware Tools"),
            ("CPU / RAM", "vCPU count and allocated RAM in GB"),
            ("Host", "ESXi host the VM is currently running on"),
            ("Snapshots", "📸 count badge — click to view/delete snapshots"),
            ("Uptime", "How long the VM has been powered on"),
            ("Actions", "Power · Snapshot · Clone · Migrate · Delete buttons"),
        ],
        col_widths=[5, 12]
    )

    heading(doc, "Available VM Actions", 2, CYAN)
    add_table(doc,
        ["Action", "Description", "Required Role"],
        [
            ("Power On / Off", "Start or shutdown the VM gracefully", "Operator, Admin"),
            ("Reboot", "Restart the guest OS", "Operator, Admin"),
            ("Clone", "Create a copy of the VM to a new name", "Operator, Admin"),
            ("Live Migrate (vMotion)", "Move running VM to another ESXi host without downtime", "Operator, Admin"),
            ("Create Snapshot", "Take a point-in-time snapshot", "Operator, Admin"),
            ("Delete Snapshot", "Remove a specific snapshot", "Operator, Admin"),
            ("Delete from Disk", "Permanently remove VM and all files", "Admin only"),
        ],
        col_widths=[5, 8, 4]
    )

    para(doc, "⚠️  WARNING: 'Delete from Disk' permanently removes the VM and all its files. This action cannot be undone.", color=RED, bold=True)
    add_page_break(doc)


def section_snapshots(doc):
    heading(doc, "7.  Snapshot Management", 1)
    divider(doc)
    para(doc,
         "The Snapshots page (nested under VMware in the sidebar) provides a centralized view of all VM snapshots "
         "across all managed vCenters.",
         color=SUB)
    add_table(doc,
        ["Column", "Description"],
        [
            ("VM Name", "The virtual machine the snapshot belongs to"),
            ("Snapshot Name", "Name given at snapshot creation"),
            ("Creation Date", "When the snapshot was taken"),
            ("Age (days)", "Number of days since creation"),
            ("vCenter", "Which vCenter environment the VM resides in"),
        ],
        col_widths=[5, 12]
    )
    para(doc, "💡 Best Practice: Aged snapshots (>7 days) consume significant datastore space and degrade VM I/O performance. Review and clean up regularly.", color=YELLOW)
    add_page_break(doc)


def section_ocp(doc):
    heading(doc, "8.  Red Hat OpenShift Module", 1)
    divider(doc)
    para(doc, "Container platform monitoring and management via the OpenShift API and Prometheus metrics.", color=SUB)
    items = [
        ("Cluster Health", "Overall cluster status, API endpoint health, and OCP version information."),
        ("Node Status", "Master/Worker/Infra node counts with Ready/NotReady breakdown."),
        ("Pod Overview", "Running, Pending, Failed pod counts. Drill into namespaces for troubleshooting."),
        ("Cluster Operators", "Available/Degraded/Progressing counts with degraded highlighting."),
        ("Resource Usage", "CPU/Memory utilisation gauges per cluster using Prometheus metrics."),
        ("Alerts", "Firing alerts from OpenShift Alertmanager surfaced in the portal."),
    ]
    for title, desc in items:
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(5)
        r1 = p.add_run(f"{title}: ")
        r1.bold = True; r1.font.color.rgb = LIGHT; r1.font.size = Pt(11)
        r2 = p.add_run(desc)
        r2.font.color.rgb = SUB; r2.font.size = Pt(11)

    para(doc,
         "Note: The portal flags a cluster as Degraded if >25% of nodes are NotReady OR >15% of Cluster Operators "
         "are degraded. Transient cycling during upgrades may briefly trigger this state.",
         color=MUTE)
    add_page_break(doc)


def section_nutanix(doc):
    heading(doc, "9.  Nutanix AHV Module", 1)
    divider(doc)
    para(doc, "The Nutanix module connects to one or more Prism Central instances.", color=SUB)
    for item in [
        "Cluster Overview — Host count, VM count, Storage capacity, AOS version",
        "VM Status — Running/stopped VM breakdown with resource allocation",
        "Storage — Storage pool usage, container free space",
        "Alerts — Critical/Warning/Info alerts from Prism Central surfaced in the portal",
        "Hardware Health — Node/disk/NIC health status from Foundation",
    ]:
        bullet(doc, item)
    add_page_break(doc)


def section_ansible(doc):
    heading(doc, "10.  Ansible AAP Module", 1)
    divider(doc)
    para(doc, "Ansible Automation Platform integration for orchestration and automation.", color=SUB)
    add_table(doc,
        ["Feature", "Description"],
        [
            ("Job Templates", "Browse all job templates defined in AAP. Launch templates directly from the portal."),
            ("Job History", "View recent job runs with status (Successful/Failed/Running), duration, and launched-by user."),
            ("Inventory", "Host inventory groups, host counts, and inventory source sync status."),
            ("Dashboard Stats", "Total hosts, recent job success rate, currently running jobs, and organisation summary."),
        ],
        col_widths=[5, 12]
    )
    add_page_break(doc)


def section_capacity(doc):
    heading(doc, "11.  Capacity Planning", 1)
    divider(doc)
    para(doc, "Resource utilisation monitoring and forecasting for CPU, RAM, and Storage.", color=SUB)
    for item in [
        "Utilisation Gauges — animated arc gauges showing used%, colour-coded: green <60%, yellow 60–80%, red >80%",
        "ESXi Host Table — per-host CPU/RAM free, VM count, and power state",
        "Datastore Table — name, type, capacity, free space, and usage bar",
        "Trend Charts — 7/14/30-day historical trend lines for CPU and RAM",
    ]:
        bullet(doc, item)
    para(doc, "⚠️  Action required when any resource exceeds 80% utilisation. Notify your infrastructure team.", color=YELLOW, bold=True)
    add_page_break(doc)


def section_roles(doc):
    heading(doc, "17.  User Roles & Permissions", 1)
    divider(doc)
    para(doc, "LaaS Portal enforces role-based access control at the API level.", color=SUB)

    add_table(doc,
        ["Role", "Description", "Key Permissions"],
        [
            ("Admin", "Full administrative access", "All operations including user management and VM deletion"),
            ("Operator", "Day-to-day operations", "VM power/clone/migrate/snapshot, launch Ansible jobs, view all reports"),
            ("Viewer", "Read-only access", "View dashboards, reports, VM list, IPAM — no write operations"),
            ("Requester", "Self-service requests", "Raise VM requests only — limited to VM Requests and Overview"),
        ],
        col_widths=[3.5, 5, 9]
    )

    heading(doc, "Permission Matrix", 2, CYAN)
    add_table(doc,
        ["Feature / Action", "Admin", "Operator", "Viewer", "Requester"],
        [
            ("View Overview Dashboard", "✅", "✅", "✅", "✅"),
            ("View VM List", "✅", "✅", "✅", "❌"),
            ("Power On/Off/Reboot VM", "✅", "✅", "❌", "❌"),
            ("Clone VM", "✅", "✅", "❌", "❌"),
            ("Migrate VM (vMotion)", "✅", "✅", "❌", "❌"),
            ("Create Snapshot", "✅", "✅", "❌", "❌"),
            ("Delete Snapshot", "✅", "✅", "❌", "❌"),
            ("Delete VM from Disk", "✅", "❌", "❌", "❌"),
            ("View Capacity Reports", "✅", "✅", "✅", "❌"),
            ("View Chargeback", "✅", "✅", "✅", "❌"),
            ("View Audit Log", "✅", "✅", "❌", "❌"),
            ("Manage Users & Roles", "✅", "❌", "❌", "❌"),
            ("Raise VM Request", "✅", "✅", "✅", "✅"),
            ("Approve/Reject VM Request", "✅", "❌", "❌", "❌"),
            ("Launch Ansible Jobs", "✅", "✅", "❌", "❌"),
            ("Export Reports (CSV/PDF)", "✅", "✅", "❌", "❌"),
        ],
        col_widths=[7, 2, 2.5, 2, 2.8]
    )
    add_page_break(doc)


def section_techspecs(doc):
    heading(doc, "18.  Technical Specifications", 1)
    divider(doc)

    heading(doc, "Application Architecture", 2, CYAN)
    para(doc, "LaaS Portal is a 3-tier web application:", color=SUB)
    for item in [
        "Frontend  →  React 18 SPA (Vite 7 build) served as static files",
        "Backend API  →  FastAPI (Python 3.11+) with SQLite database",
        "Data Sources  →  VMware vCenter REST API · OCP API + Prometheus · Nutanix Prism Central API · Ansible AAP API · phpIPAM REST API",
    ]:
        bullet(doc, item)

    heading(doc, "Server Hardware Specifications", 2, CYAN)
    add_table(doc,
        ["Component", "Specification"],
        [
            ("CPU", "Intel Xeon Gold 6130 — 2× 8-core / 16-thread = 16 physical cores / 32 threads"),
            ("RAM", "24 GB DDR4 ECC (usable ~22 GB)"),
            ("Storage", "300 GB SSD — Application, OS, database"),
            ("Network", "1 Gbps NIC to corporate LAN"),
            ("Operating System", "Windows Server 2019 Standard"),
            ("Expected Concurrent Users", "30–60 (up to 100+ with software optimisations)"),
        ],
        col_widths=[6, 12]
    )

    heading(doc, "Software Stack", 2, CYAN)
    add_table(doc,
        ["Component", "Technology / Version"],
        [
            ("Frontend Framework", "React 18 + Vite 7"),
            ("UI Library", "Custom CSS (no external UI framework)"),
            ("Backend Framework", "FastAPI 0.111+ (Python 3.11+)"),
            ("ASGI Server", "Uvicorn (recommended: 4–8 workers)"),
            ("Database", "SQLite 3 with WAL mode enabled"),
            ("VMware SDK", "pyVmomi (vSphere SDK for Python)"),
            ("OCP Integration", "kubernetes Python client + Prometheus HTTP API"),
            ("Nutanix Integration", "Prism Central REST API v3"),
            ("Ansible Integration", "Ansible AAP REST API v2"),
            ("IPAM Integration", "phpIPAM REST API"),
        ],
        col_widths=[6, 12]
    )

    heading(doc, "Key API Endpoints", 2, CYAN)
    add_table(doc,
        ["Method", "Endpoint", "Description"],
        [
            ("GET", "/api/vcenters", "List all registered vCenter instances"),
            ("GET", "/api/vms/{vc_id}", "Fetch all VMs for a vCenter"),
            ("POST", "/api/vm/{vc_id}/{vm_name}/power", "Power on/off/reboot a VM"),
            ("POST", "/api/vm/{vc_id}/{vm_name}/snapshot", "Create a VM snapshot"),
            ("DELETE", "/api/vm/{vc_id}/{vm_name}/snapshot", "Delete a VM snapshot"),
            ("POST", "/api/vm/{vc_id}/{vm_name}/clone", "Clone a VM"),
            ("POST", "/api/vm/{vc_id}/{vm_name}/migrate", "vMotion a VM to another host"),
            ("GET", "/api/ocp/clusters", "List all OCP clusters"),
            ("GET", "/api/nutanix/pcs", "List all Nutanix Prism Centrals"),
            ("GET", "/api/audit", "Fetch audit log entries"),
        ],
        col_widths=[2.5, 7.5, 8]
    )
    add_page_break(doc)


def section_faq(doc):
    heading(doc, "19.  FAQ & Troubleshooting", 1)
    divider(doc)
    faq_items = [
        ("The portal shows '0 VMs' but vCenter has VMs",
         "Check that the vCenter is registered and credentials are valid. Try clicking Retry. If the issue persists, check the backend logs for vCenter API errors."),
        ("I cannot see the Audit Log or Users menu",
         "Audit Log requires Operator or Admin role. Users & Roles requires Admin role. Contact your Admin to upgrade your role."),
        ("Snapshot creation timed out",
         "Large VMs on busy datastores can exceed the portal timeout. Check the vCenter task list directly. Retry during off-peak hours."),
        ("OCP cluster shows 'Degraded' but seems healthy",
         "The portal flags Degraded if >25% of nodes are NotReady OR >15% of Cluster Operators are degraded. Transient cycling during upgrades can trigger this briefly."),
        ("How do I add a new vCenter / OCP cluster / Nutanix?",
         "Navigate to Settings (Admin only) → Add vCenter / Add OCP Cluster / Add Nutanix PC. Enter the endpoint URL and credentials, then save."),
        ("Can I export VM lists or capacity reports?",
         "Yes — most tables have a CSV export (Admin and Operator). Chargeback also has PDF export. Audit log export requires Admin."),
        ("The analog clock shows the wrong time",
         "The clock uses the browser's local system clock. Ensure your workstation is synchronized to the corporate NTP server."),
    ]
    for q, a in faq_items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f"Q: {q}")
        r.bold = True; r.font.color.rgb = ACCENT; r.font.size = Pt(12)
        para(doc, f"A: {a}", color=SUB, indent=0.5)
        divider(doc)

    add_page_break(doc)


def section_auth(doc):
    heading(doc, "20.  Document Authorization", 1)
    divider(doc)

    table = doc.add_table(rows=2, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    data = [
        ("✍️  Prepared By", "Sekhar Perumal", "SDX Infrastructure & DC Operations · Wipro", "March 2026", ACCENT),
        ("✅  Approved By", "Khalid Khan",    "SDX Infrastructure & DC Operations · Wipro", "March 2026", GREEN),
    ]
    for col_idx, (role, name, dept, date, color) in enumerate(data):
        cell = table.rows[0].cells[col_idx]
        shade_cell(cell, DARK_BG)
        p = cell.paragraphs[0]
        r = p.add_run(role)
        r.font.size = Pt(10); r.font.color.rgb = MUTE; r.bold = True

        cell2 = table.rows[1].cells[col_idx]
        shade_cell(cell2, DARK_BG)
        p2 = cell2.paragraphs[0]
        r1 = p2.add_run(name + "\n")
        r1.font.size = Pt(15); r1.bold = True; r1.font.color.rgb = LIGHT
        r2 = p2.add_run(dept + "\n")
        r2.font.size = Pt(10); r2.font.color.rgb = color
        r3 = p2.add_run(date)
        r3.font.size = Pt(10); r3.font.color.rgb = MUTE

    set_col_width(table, 0, 9)
    set_col_width(table, 1, 9)

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("LaaS Portal — End User Manual v6.0   ·   Wipro SDX INFRA & DC Operations   ·   March 2026")
    r.font.size = Pt(10); r.font.color.rgb = MUTE
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("Confidential — Internal Use Only")
    r2.font.size = Pt(10); r2.font.color.rgb = MUTE; r2.italic = True


# ── MAIN ─────────────────────────────────────────────────────────────────────
def build():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.page_width  = Cm(21)
        section.page_height = Cm(29.7)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2)
        section.bottom_margin = Cm(2)

    # Document-wide font default
    doc.styles['Normal'].font.name = 'Segoe UI'
    doc.styles['Normal'].font.size = Pt(11)
    doc.styles['Normal'].font.color.rgb = SUB

    # Dark-mode document background
    doc.styles['Normal'].paragraph_format.space_after = Pt(6)

    cover_page(doc)
    toc_page(doc)
    section_intro(doc)
    section_sysreq(doc)
    section_quickstart(doc)
    section_navigation(doc)
    section_overview(doc)
    section_vmware(doc)
    section_snapshots(doc)
    section_ocp(doc)
    section_nutanix(doc)
    section_ansible(doc)
    section_capacity(doc)

    # Sections 12–16 (summarised)
    heading(doc, "12.  Project Utilization", 1)
    divider(doc)
    para(doc, "Displays resource consumption broken down by project (cost center / tag). Each project row shows VM count, vCPU, RAM (GB), Storage (GB), running vs stopped VM ratio, and a visual bar showing the project's share of total resources.", color=SUB)
    add_page_break(doc)

    heading(doc, "13.  Chargeback", 1)
    divider(doc)
    para(doc, "Calculates monthly infrastructure costs per project based on configurable rate cards (vCPU/month, RAM GB/month, Storage GB/month). Features: month selector, project filter, export to PDF/CSV, donut chart breakdown, total cost summary.", color=SUB)
    add_page_break(doc)

    heading(doc, "14.  Networks & IPAM", 1)
    divider(doc)
    for item in [
        "Networks Page — vCenter port groups and distributed switches with VLAN IDs and connected VM counts.",
        "IPAM Page — Subnet overview with used/free IP counts, occupancy % bar, IP allocation table.",
        "Free IP Finder — Search for available IPs within a subnet before provisioning new VMs.",
        "Allocation History — Track when IPs were allocated and to which VM/host.",
    ]:
        bullet(doc, item)
    add_page_break(doc)

    heading(doc, "15.  VM Request Workflow", 1)
    divider(doc)
    for i, (title, body) in enumerate([
        ("Requester Raises a Request", "Navigate to VM Requests → New Request. Fill in VM Name, vCPU, RAM, Storage, OS, Project, justification, and submit."),
        ("Admin Receives Notification", "Admin sees the pending badge counter on the VM Requests nav item."),
        ("Admin Reviews & Decides", "Admin opens the request, reviews specs, then clicks Approve or Reject with optional comments."),
        ("VM Provisioned or Rejected", "If approved, VM is provisioned from a template. If rejected, the requester sees the reason."),
    ], 1):
        p = doc.add_paragraph()
        r1 = p.add_run(f"Step {i}: {title}  — ")
        r1.bold = True; r1.font.color.rgb = ACCENT; r1.font.size = Pt(11)
        r2 = p.add_run(body)
        r2.font.color.rgb = SUB; r2.font.size = Pt(11)
    add_page_break(doc)

    heading(doc, "16.  Audit Log", 1)
    divider(doc)
    para(doc, "Every user action is recorded with: Timestamp · User · Action Type · Resource · Status · Details.", color=SUB)
    para(doc, "Filters: Date range · User · Action type · Status · Free text search. Export: CSV (Admin only).", color=SUB)
    add_page_break(doc)

    section_roles(doc)
    section_techspecs(doc)
    section_faq(doc)
    section_auth(doc)

    doc.save(OUT)
    print(f"✅  Saved: {OUT}")


if __name__ == "__main__":
    build()
