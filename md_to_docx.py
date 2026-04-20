import re
from pathlib import Path

SRC = Path(r"c:/caas-dashboard/END_USER_TECHNICAL_DOCUMENTATION.md")
OUT = Path(r"c:/caas-dashboard/END_USER_TECHNICAL_DOCUMENTATION.docx")

from docx import Document

doc = Document()

def clean_inline(text: str) -> str:
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"\*([^*]+)\*", r"\1", text)
    return text

in_code = False
for raw in SRC.read_text(encoding="utf-8").splitlines():
    line = raw.rstrip("\n")

    if line.strip().startswith("```"):
        in_code = not in_code
        continue

    if in_code:
        p = doc.add_paragraph(clean_inline(line))
        if p.runs:
            p.runs[0].font.name = "Consolas"
        continue

    if not line.strip():
        doc.add_paragraph("")
        continue

    if line.startswith("### "):
        doc.add_heading(clean_inline(line[4:].strip()), level=3)
        continue
    if line.startswith("## "):
        doc.add_heading(clean_inline(line[3:].strip()), level=2)
        continue
    if line.startswith("# "):
        doc.add_heading(clean_inline(line[2:].strip()), level=1)
        continue

    m = re.match(r"^\s*[-*]\s+(.*)$", line)
    if m:
        doc.add_paragraph(clean_inline(m.group(1)), style="List Bullet")
        continue

    m = re.match(r"^\s*\d+[\.)]\s+(.*)$", line)
    if m:
        doc.add_paragraph(clean_inline(m.group(1)), style="List Number")
        continue

    if "|" in line and line.strip().startswith("|"):
        txt = "  ".join([c.strip() for c in line.strip().strip("|").split("|")])
        doc.add_paragraph(clean_inline(txt))
        continue

    doc.add_paragraph(clean_inline(line))

doc.save(OUT)
print(f"CREATED: {OUT}")
